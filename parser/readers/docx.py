"""
DOCX file reader implementation using python-docx.

Provides text and metadata extraction from DOCX resume files with
robust error handling and support for various formatting elements.
"""

import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.xmlchemy import OxmlElement

from parser.exceptions import FileReadException

logger = logging.getLogger(__name__)


class DocxReader:
    """Reader for DOCX resume files using python-docx."""

    def __init__(self, file_path: str | Path) -> None:
        """
        Initialize DOCX reader.

        Args:
            file_path: Path to the DOCX file.
        """
        self.file_path = Path(file_path)

    def extract_text(self) -> str:
        """
        Extract text from DOCX file.

        Handles paragraphs, tables, and common formatting while preserving
        structure through line breaks.

        Returns:
            Extracted text content.

        Raises:
            FileReadException: If DOCX reading fails.
        """
        try:
            doc = Document(self.file_path)
            text_content = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            extracted_text = "\n".join(text_content)
            if not extracted_text.strip():
                logger.warning(f"No text extracted from DOCX: {self.file_path}")

            logger.debug(f"Extracted {len(extracted_text)} characters from DOCX")
            return extracted_text

        except Exception as e:
            msg = f"Failed to read DOCX file {self.file_path}: {e}"
            logger.error(msg)
            raise FileReadException(msg) from e

    def extract_metadata(self) -> dict[str, Any]:
        """
        Extract metadata from DOCX file.

        Returns:
            Dictionary with document properties and statistics.

        Raises:
            FileReadException: If metadata extraction fails.
        """
        try:
            doc = Document(self.file_path)
            core_props = doc.core_properties

            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
            }

            logger.debug(f"Extracted metadata from DOCX: {metadata}")
            return metadata

        except Exception as e:
            msg = f"Failed to extract metadata from DOCX {self.file_path}: {e}"
            logger.error(msg)
            raise FileReadException(msg) from e
